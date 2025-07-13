import os
import sys
import re
from PyQt5.QtWidgets import (
    QApplication, QWidget, QHBoxLayout, QVBoxLayout, QSizePolicy, QFileDialog,
    QDialog, QLabel, QScrollArea, QLineEdit, QShortcut, QPushButton, QMessageBox
)
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QKeySequence, QPixmap, QIcon
from PyQt5.QtWidgets import QSplashScreen
from docx.enum.text import WD_TAB_ALIGNMENT

from modules.sidebar import create_sidebar
from modules.header import create_header
from modules.filters import create_filter_section
from modules.searchbar import create_search_bar
from modules.results_box import ResultBoxWidget
from modules.data_table import create_data_table, get_data_table, set_table_model
from modules.status_section import StatusSection
from modules.footer_credit import FooterCredit
from docx.shared import Inches, Pt

from functions.display_table import load_data, filter_data, get_loaded_data, valid_filter_columns
from functions.pandas_table_model import PandasTableModel
from functions.filter_workers import FilterWorker
from functions.resume_dialog import ResumeDialog


class GpaidiaApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Aplikasi Manajemen SIGADISIndo")
        self.setWindowIcon(QIcon(resource_path("icons/splash.png")))
        
        # PERBAIKAN: Set ukuran minimum window yang cukup besar
        self.setMinimumSize(1200, 800)
        self.setGeometry(100, 100, 1400, 900)  # Ukuran default lebih besar
        
        self.is_sidebar_mini = False
        self.result_box_widget = None
        self.search_input = None
        self.filtered_data = None
        self.worker_thread = None
        self.status_section = None

        self.filter_widget = create_filter_section(available_filters={})

        self.search_timer = QTimer()
        self.search_timer.setSingleShot(True)
        self.search_timer.timeout.connect(self.apply_search_filter)

        self.setup_ui()
        self.setStyleSheet("background-color: white;")

    def setup_ui(self):
        # PERBAIKAN: Layout utama dengan margin dan spacing yang konsisten
        self.main_layout = QHBoxLayout(self)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)

        # Setup filter widget
        self.filter_widget = create_filter_section(available_filters={})
        self.filter_widget.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Maximum)
        self.filter_widget.setMinimumWidth(450)
        self.filter_widget.filter_changed.connect(self.apply_filter)

        # PERBAIKAN: Sidebar dengan size policy yang tepat
        self.sidebar_frame, self.sidebar_state_handler = create_sidebar(
            on_load_callback=self.handle_file_loaded,
            filter_widget=self.filter_widget,
            on_reset_callback=self.reset_filters
        )
        # Set size policy untuk sidebar agar tidak mengambil ruang berlebihan
        self.sidebar_frame.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Expanding)
        self.main_layout.addWidget(self.sidebar_frame)

        # PERBAIKAN: Content area dengan layout yang lebih stabil
        self.content_layout = QVBoxLayout()
        self.content_layout.setContentsMargins(0, 0, 0, 0)
        self.content_layout.setSpacing(0)

        # Header
        header_widget = create_header(toggle_callback=self.toggle_sidebar)
        self.content_layout.addWidget(header_widget)

        # PERBAIKAN: Filter dan result box dengan layout yang lebih stabil
        self.filter_result_layout = QHBoxLayout()
        self.filter_result_layout.setContentsMargins(20, 10, 20, 10)
        self.filter_result_layout.setSpacing(30)

        self.result_box_widget = ResultBoxWidget()
        self.result_box_widget.resume_clicked.connect(self.show_resume_dialog)

        # PERBAIKAN: Gunakan stretch ratio yang lebih seimbang
        self.filter_result_layout.addWidget(self.filter_widget, stretch=2, alignment=Qt.AlignTop)
        self.filter_result_layout.addWidget(self.result_box_widget, stretch=1, alignment=Qt.AlignTop)

        # PERBAIKAN: Container untuk filter dan result dengan size policy yang tepat
        filter_result_container = QWidget()
        filter_result_container.setLayout(self.filter_result_layout)
        filter_result_container.setStyleSheet("background-color: white;")
        filter_result_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Maximum)
        self.content_layout.addWidget(filter_result_container)

        # Search bar
        search_widget, self.search_input = create_search_bar(self.search_data)
        self.content_layout.addWidget(search_widget)

        # PERBAIKAN: Table dan status section dengan wrapper yang lebih stabil
        table_container = create_data_table()
        table_view = get_data_table()
        table_view.doubleClicked.connect(self.show_row_preview)

        # PERBAIKAN: Wrapper untuk table dan status dengan size policy yang tepat
        table_and_status_wrapper = QWidget()
        table_and_status_layout = QVBoxLayout()
        table_and_status_layout.setContentsMargins(0, 0, 0, 5)
        table_and_status_layout.setSpacing(5)

        # PERBAIKAN: Table container dengan size policy yang tepat
        table_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        table_and_status_layout.addWidget(table_container)

        # Status section
        self.status_section = StatusSection()
        self.status_section.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        table_and_status_layout.addWidget(self.status_section)

        # PERBAIKAN: Wrapper dengan size policy yang tepat
        table_and_status_wrapper.setLayout(table_and_status_layout)
        table_and_status_wrapper.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.content_layout.addWidget(table_and_status_wrapper)

        # PERBAIKAN: Footer dengan size policy yang tepat
        footer = FooterCredit()
        footer.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.content_layout.addWidget(footer)

        # PERBAIKAN: Tambahkan stretch di akhir untuk mendorong semua ke atas
        self.content_layout.addStretch(1)

        # Selection model setup
        if table_view.selectionModel():
            table_view.selectionModel().selectionChanged.connect(self.on_selection_changed)

        # PERBAIKAN: Content widget dengan size policy yang tepat
        content_widget = QWidget()
        content_widget.setLayout(self.content_layout)
        content_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.main_layout.addWidget(content_widget)

        # Shortcut setup
        shortcut = QShortcut(QKeySequence(Qt.Key_Escape), self)
        shortcut.activated.connect(lambda: get_data_table().clearSelection())

    def show_resume_dialog(self):
        try:
            # Ambil filter aktif dari FilterWidget
            filters = self.filter_widget.get_current_filters()
            
            # Ambil nilai-nilai box hasil
            counts = {
                label: self.result_box_widget.box_labels[label].text()
                for label in ["Jumlah PNS", "Jumlah NON PNS", "Jumlah PPPK", "Jumlah Sekolah"]
            }

            search_text = self.search_input.text().strip() if self.search_input else ""
            dialog = ResumeDialog(self.result_box_widget.font_family, filters, counts, search_text, self)

            dialog.exec_()
        except Exception as e:
            print(f"❌ Gagal tampilkan resume popup: {e}")

    def toggle_sidebar(self):
        self.is_sidebar_mini = not self.is_sidebar_mini
        self.sidebar_state_handler(self.is_sidebar_mini)
        
        # PERBAIKAN: Force update layout setelah toggle
        self.main_layout.update()
        self.updateGeometry()

    def handle_file_loaded(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Pilih File Data", "", "Excel Files (*.xlsx *.xls);;CSV Files (*.csv)"
        )
        if path:
            load_data(path)
            
            # PERBAIKAN: Lebih careful dalam mengganti filter widget
            old_filter = self.filter_widget
            self.filter_result_layout.removeWidget(old_filter)
            old_filter.setParent(None)
            old_filter.deleteLater()  # Pastikan widget lama dihapus

            self.filter_widget = create_filter_section(available_filters=valid_filter_columns)
            self.filter_widget.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Maximum)
            self.filter_widget.filter_changed.connect(self.apply_filter)
            self.filter_widget.setMinimumWidth(450)
            
            # PERBAIKAN: Insert dengan stretch yang konsisten
            self.filter_result_layout.insertWidget(0, self.filter_widget, stretch=2, alignment=Qt.AlignTop)

            self.reset_filters()

    def reset_filters(self):
        if self.filter_widget:
            self.filter_widget.reset_filters()

        if self.search_input:
            self.search_input.setText("")

        get_data_table().clearSelection()

        if self.status_section:
            self.status_section.update_selected_row(None)

        df = get_loaded_data()
        self.filtered_data = df
        self.apply_search_filter()

    def apply_filter(self, filters):
        df = filter_data(filters)
        self.filtered_data = df
        self.apply_search_filter()

    def apply_search_filter(self):
        df = self.filtered_data if self.filtered_data is not None else get_loaded_data()
        if df is None:
            return

        query = self.search_input.text().strip().lower() if self.search_input else ""

        if self.worker_thread:
            self.worker_thread.quit()
            self.worker_thread.wait()

        self.worker_thread = FilterWorker(df, query)
        self.worker_thread.result_ready.connect(self.on_filter_result)
        self.worker_thread.start()

    def on_filter_result(self, result_df):
        model = PandasTableModel(result_df)
        set_table_model(model)

        if self.result_box_widget:
            self.result_box_widget.update_counts(result_df)

        if self.status_section:
            self.status_section.update_total(len(result_df))

        table_view = get_data_table()
        sel_model = table_view.selectionModel()

        try:
            sel_model.selectionChanged.disconnect(self.on_selection_changed)
        except (TypeError, RuntimeError):
            pass

        sel_model.selectionChanged.connect(self.on_selection_changed)

    def on_selection_changed(self):
        table_view = get_data_table()
        sel_model = table_view.selectionModel()
        selected_rows = list(set(index.row() for index in sel_model.selectedIndexes()))

        if selected_rows:
            row_index = selected_rows[0]
            if self.status_section:
                self.status_section.update_selected_row(row_index)
        else:
            if self.status_section:
                self.status_section.update_selected_row(None)

    def search_data(self, text):
        self.search_timer.start(300)

    # PERBAIKAN: Tambahkan method untuk handle resize event
    def resizeEvent(self, event):
        super().resizeEvent(event)
        # Update layout ketika window diresize
        self.main_layout.update()

    def show_row_preview(self, index):
        table_view = get_data_table()
        model = index.model()

        visible_columns = [
            col for col in range(model.columnCount())
            if not table_view.isColumnHidden(col)
        ]

        dialog = QDialog(self)
        dialog.setWindowTitle("Preview Detail")
        dialog.setMinimumSize(1100, 800)
        dialog.setStyleSheet("background-color: white;")

        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(20)

        title = QLabel("Preview")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 20px; font-weight: bold; color: #212121;")
        main_layout.addWidget(title)

        search_input = QLineEdit()
        search_input.setPlaceholderText("Cari kolom atau isi...")
        search_input.setStyleSheet("""
            QLineEdit {
                padding: 8px;
                font-size: 14px;
                border: 1px solid #ccc;
                border-radius: 6px;
            }
        """)
        main_layout.addWidget(search_input)

        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)

        scroll_widget = QWidget()
        scroll_layout = QHBoxLayout(scroll_widget)
        scroll_layout.setSpacing(30)

        col1, col2, col3, col4 = QVBoxLayout(), QVBoxLayout(), QVBoxLayout(), QVBoxLayout()
        col_widgets = []

        part = (len(visible_columns) + 3) // 4

        for idx, col in enumerate(visible_columns):
            col_name = model.headerData(col, Qt.Horizontal)
            value = model.data(model.index(index.row(), col), Qt.DisplayRole)

            label = QLabel(f"<b>{col_name}</b>: {value}")
            label.setWordWrap(True)
            label.setTextInteractionFlags(Qt.TextSelectableByMouse)
            label.setStyleSheet("font-size: 13px; color: #333;")
            col_widgets.append((label, col_name.lower(), str(value).lower()))

            if idx < part:
                col1.addWidget(label)
            elif idx < 2 * part:
                col2.addWidget(label)
            elif idx < 3 * part:
                col3.addWidget(label)
            else:
                col4.addWidget(label)

        scroll_layout.addLayout(col1)
        scroll_layout.addLayout(col2)
        scroll_layout.addLayout(col3)
        scroll_layout.addLayout(col4)

        scroll_area.setWidget(scroll_widget)
        main_layout.addWidget(scroll_area)

        def highlight_search():
            query = search_input.text().lower().strip()
            found = False
            for label, col_name, val in col_widgets:
                combined = f"{col_name} {val}"
                if query and query in combined:
                    label.setStyleSheet("font-size: 13px; color: black; background-color: #C1A910; padding: 2px;")
                    if not found:
                        scroll_area.ensureWidgetVisible(label)
                        found = True
                else:
                    label.setStyleSheet("font-size: 13px; color: #333;")

        search_input.textChanged.connect(highlight_search)

        export_layout = QHBoxLayout()
        export_layout.setSpacing(15)

        export_word_btn = QPushButton("Export To Word")
        export_excel_btn = QPushButton("Export To Excel")

        export_word_btn.setStyleSheet("background-color: #2B579A; color: white; font-size: 22px; padding: 20px 40px; border-radius: 5px;")
        export_excel_btn.setStyleSheet("background-color: #10793F; color: white; font-size: 22px; padding: 20px 40px; border-radius: 5px;")

        export_layout.addWidget(export_word_btn)
        export_layout.addWidget(export_excel_btn)
        main_layout.addLayout(export_layout)

        def export_to_word():
            from docx import Document
            file_path, _ = QFileDialog.getSaveFileName(dialog, "Simpan sebagai Word", "", "Word Files (*.docx)")
            if file_path:
                doc = Document()
                doc.add_heading('Preview Data Pegawai', level=1)
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'

                for label, _, _ in col_widgets:
                    clean_text = re.sub(r"<[^>]+>", "", label.text())
                    kolom, nilai = clean_text.split(": ", 1) if ": " in clean_text else (clean_text, "")
                    row_cells = table.add_row().cells
                    row_cells[0].text = kolom.strip()
                    row_cells[1].text = nilai.strip()
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(11)
                doc.save(file_path)
                QMessageBox.information(None, "Berhasil", "Data berhasil disimpan dalam format Word.")

        def export_to_excel():
            import pandas as pd
            file_path, _ = QFileDialog.getSaveFileName(dialog, "Simpan sebagai Excel", "", "Excel Files (*.xlsx)")
            if file_path:
                data_dict = {col_name.title(): [val] for _, col_name, val in col_widgets}
                df = pd.DataFrame(data_dict)
                df.to_excel(file_path, index=False)
                QMessageBox.information(None, "Berhasil", "Data berhasil disimpan dalam format Excel.")

        export_word_btn.clicked.connect(export_to_word)
        export_excel_btn.clicked.connect(export_to_excel)

        dialog.setLayout(main_layout)
        dialog.exec_()

def resource_path(relative_path):
    base_path = getattr(sys, '_MEIPASS', os.path.abspath("."))
    return os.path.join(base_path, relative_path)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    splash_path = resource_path("icons/splash.png")
    splash_pix = QPixmap(splash_path)
    splash = QSplashScreen(splash_pix, Qt.WindowStaysOnTopHint)
    splash.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
    splash.show()

    window = GpaidiaApp()

    def start_app():
        splash.finish(window)
        window.show()

    QTimer.singleShot(1, start_app)
    sys.exit(app.exec_())