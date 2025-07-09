from PyQt5.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QFileDialog, QMessageBox
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ResumeDialog(QDialog):
    def __init__(self, font_family, filters, counts, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Ringkasan Data")
        self.setMinimumWidth(400)
        self.font_family = font_family
        self.filters = filters
        self.counts = counts

        layout = QVBoxLayout()

        title = QLabel("🧾 Ringkasan Data")
        title.setFont(QFont(font_family, 12, QFont.Bold))
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        layout.addSpacing(10)

        self.summary_texts = []

        if filters:
            filter_title = QLabel("📌 Filter Aktif:")
            filter_title.setFont(QFont(font_family, 10, QFont.Bold))
            layout.addWidget(filter_title)
            self.summary_texts.append("📌 Filter Aktif:")

            for key, value in filters.items():
                line = f"- {key}: {value}"
                lbl = QLabel(line)
                lbl.setFont(QFont(font_family, 10))
                layout.addWidget(lbl)
                self.summary_texts.append(line)
        else:
            lbl = QLabel("Tidak ada filter aktif.")
            layout.addWidget(lbl)
            self.summary_texts.append("Tidak ada filter aktif.")

        layout.addSpacing(10)

        count_title = QLabel("📊 Statistik Data:")
        count_title.setFont(QFont(font_family, 10, QFont.Bold))
        layout.addWidget(count_title)
        self.summary_texts.append("📊 Statistik Data:")

        for key, value in counts.items():
            line = f"- {key}: {value}"
            lbl = QLabel(line)
            lbl.setFont(QFont(font_family, 10))
            layout.addWidget(lbl)
            self.summary_texts.append(line)

        layout.addSpacing(15)

        # Tombol Save to Word (tanpa tombol OK)
        save_button = QPushButton("💾 Save to Word")
        save_button.setFont(QFont(font_family, 10, QFont.Bold))
        save_button.setCursor(Qt.PointingHandCursor)
        save_button.setStyleSheet("""
            QPushButton {
                background-color: #2B579A;   /* Warna khas Word */
                color: white;
                border-radius: 6px;
                padding: 8px 16px;
            }
            QPushButton:hover {
                background-color: #1E3E73;
            }
        """)
        save_button.clicked.connect(self.save_to_word)
        layout.addWidget(save_button, alignment=Qt.AlignCenter)

        self.setLayout(layout)

    def save_to_word(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Simpan Ringkasan", "ringkasan.docx", "Word Document (*.docx)")
        if file_path:
            try:
                doc = Document()

                # Judul di tengah
                title = doc.add_paragraph()
                run = title.add_run("Ringkasan Data")
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 102, 204)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph("")

                for line in self.summary_texts:
                    doc.add_paragraph(line)

                doc.save(file_path)

                # ✅ Notifikasi Berhasil
                QMessageBox.information(self, "Berhasil", "File Word berhasil disimpan!")

            except Exception as e:
                # ❌ Notifikasi Gagal
                QMessageBox.critical(self, "Gagal", f"Gagal menyimpan file Word.\n\nError:\n{str(e)}")

